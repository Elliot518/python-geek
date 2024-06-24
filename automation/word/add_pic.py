import docx

doc = docx.Document()
# doc.add_picture('chimera.jpeg', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
doc.add_picture('chimera.jpeg', width=docx.shared.Cm(3), height=docx.shared.Cm(4))

doc.save('pic.docx')
print("doc saved")
