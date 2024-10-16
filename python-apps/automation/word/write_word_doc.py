import docx

doc = docx.Document()

doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
header1 = doc.add_heading('Header 4', 4)

# add a new line after header
header1.add_run().add_break()

doc.add_paragraph('Hello, world!')
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.save('helloworld.docx')

print("doc saved")
